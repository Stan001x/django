import mimetypes
from os.path import getsize
from wsgiref.util import FileWrapper
import num2rus

from django.contrib.auth.mixins import LoginRequiredMixin
from django.db.models import Q
from django.contrib.postgres.search import SearchVector, SearchQuery, SearchRank
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.shortcuts import render, redirect, reverse, get_object_or_404
from django.http import HttpResponseNotFound, JsonResponse, HttpResponse, StreamingHttpResponse
from django.contrib.auth.decorators import login_required
from django.views.generic import ListView, DetailView
from django.views import View
from django.views.generic.edit import DeleteView, FormView, CreateView, UpdateView
from django.urls import reverse_lazy
from django.core.paginator import Paginator
from django.conf import settings
from django.views.decorators.csrf import csrf_exempt
import json
from docx import Document
from docx.shared import Pt, Mm
from docxtpl import DocxTemplate, InlineImage
from django.forms import formset_factory

from .forms import UpdateReport, CreateReport, CreatePurposeOfAssessment, CreatePersonDataForm, ObjectOfAssessmentForm, AnaloguesForm, ImagesForm, AdjustmentsForm
from .models import Report, PurposeOfAssessment, Analogues, ObjectOfAssessment, ClientPersonData, Adjustments
from django import forms



def MakeReport(report):
    print('make report start')
    # document = Document('notarius/documents/notar.docx')
    # style = document.styles['Normal']
    # # название шрифта
    # style.font.name = 'Times New Roman'
    # # размер шрифта
    # style.font.size = Pt(12)
    # # Таблица с Указанием Заказчика
    # table = document.add_table(1, 2)
    # table.style = 'Table Grid'
    # # Получаем строку с колонками из добавленной таблицы
    # head_cells = table.rows[0].cells
    # # добавляем названия колонок
    # for i, item in enumerate(['Заказчик', {report.clientname}]):
    #     p = head_cells[i].paragraphs[0]
    #     p.add_run(item).bold = True
    #     p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # document.add_paragraph(f'Текст документа\n1')
    # document.save(f'notarius/documents/{report.contractNumber}.docx')
    doc = DocxTemplate('notarius/documents/notar.docx')
#    image = InlineImage(doc, image_descriptor=report.analogue1.analogueImage1.imageFile, width=Mm(90))
    context = {'report': report}
    doc.render(context)
    doc.save(f'media/documents/notarius/{report.pk}.docx')

# def MakeReport(item):
#     document = Document('notarius/documents/notar.docx')
#     style = document.styles['Normal']
#     # название шрифта
#     style.font.name = 'Times New Roman'
#     # размер шрифта
#     style.font.size = Pt(12)
#     # Таблица с Указанием Заказчика
#     table = document.add_table(1, 2)
#     table.style = 'Table Grid'
#     # Получаем строку с колонками из добавленной таблицы
#     head_cells = table.rows[0].cells
#     # добавляем названия колонок
#     for i, item in enumerate(['Кол-во', {item.clientname}]):
#         p = head_cells[i].paragraphs[0]
#     document.add_paragraph('Текст документа')
#     i = 0
#     while i < len(document.paragraphs):
#         print(f'{i} {document.paragraphs[i].text}\n')
#         i = i + 1
#     document.save(f'notarius/documents/{item.contractNumber}.docx')


# def index(request):
#   page_obj = items = Report.objects.all()
#   item_name = request.Get.get('search')
#   if item_name != '' and item_name is not None:
#       page_obj = items.filter(name__icontains=item_name)
#
#   paginator = Paginator(page_obj, 2)
#   page_number = request.GET.get('page')
#   page_obj = paginator.get_page(page_number)
#    context = {'items':items, 'page_obj':page_obj}
#    return render(request, "notarius/notarius-main.html", context)

class ReportListView(ListView):
    model = Report
    template_name = "notarius/notarius-main.html"
    context_object_name = 'items'
    paginate_by = 20

    def get_queryset(self):
        query = self.request.GET.get('search')
        if query:
            object_list = self.model.objects.filter(Q(clientName__icontains=query) | Q(reportNumber__icontains=query))
        else:
            object_list = self.model.objects.all()
        return object_list

# def indexItem(request, my_id):
#    item = Report.objects.get(id=my_id)
#    context = {
#        'item':item

#    }
#    return render(request, "notarius/report.html", context=context)

class ReportDetailView(DetailView):
    model = Report
    template_name = "notarius/report.html"
    context_object_name = 'item'
    pk_url_kwarg = 'pk'

    def get_context_data(self, **kwargs):
        context = super(ReportDetailView, self).get_context_data(**kwargs)

        context["stripe_publishable_key"] = settings.STRIPE_PUBLISHABLE_KEY
        return context


#@login_required
class AddItem(LoginRequiredMixin, CreateView):
    template_name = 'notarius/additem.html'
    form_class = CreateReport
#    model = Report
#    fields = '__all__'
    context_object_name = 'item'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        # context["purpose_of_assessment"] = CreatePurposeOfAssessment()
        context["client_person_data_form"] = CreatePersonDataForm()
        context["object_of_assessment_form"] = ObjectOfAssessmentForm()
        context["analogues_form1"] = AnaloguesForm(prefix='analogue1')
        context["analogues_form2"] = AnaloguesForm(prefix='analogue2')
        context["analogues_form3"] = AnaloguesForm(prefix='analogue3')
        # context["analogues_image11_form"] = ImagesForm(prefix='analogues_image11')
        # context["analogues_image21_form"] = ImagesForm(prefix='analogues_image21')
        # context["analogues_image31_form"] = ImagesForm(prefix='analogues_image31')
        context["adjustment_form1"] = AdjustmentsForm(prefix='analogue1')
        context["adjustment_form2"] = AdjustmentsForm(prefix='analogue2')
        context["adjustment_form3"] = AdjustmentsForm(prefix='analogue3')
        # if self.request.POST.get('clientType') == '2':
        #      print('лицо')
        #      context["client_person_data_form"] = None
             # return render(self.request, "notarius/test.html", context={"client_person_data_form": '1', })

        # self.purposeOfAssessment_id = self.request.POST.get('purposeOfAssessment')
        # if self.request.method == "POST":
        #     PurposeOfAssessment.objects.create(purposeOfAssessment1=self.request.POST.get('purposeOfAssessment1'))
        return context
#    success_url = reverse_lazy("notarius:index")


    def post(self, request, *args, **kwargs):
        form = self.form_class(request.POST, request.FILES)
        client_person_data_form = CreatePersonDataForm(request.POST)
        object_of_assessment_form = ObjectOfAssessmentForm(request.POST)
        analogues_form1 = AnaloguesForm(request.POST, prefix='analogue1')
        analogues_form2 = AnaloguesForm(request.POST, prefix='analogue2')
        analogues_form3 = AnaloguesForm(request.POST, prefix='analogue3')
        adjustment_form1 = AdjustmentsForm(request.POST, prefix='analogue1')
        adjustment_form2 = AdjustmentsForm(request.POST, prefix='analogue2')
        adjustment_form3 = AdjustmentsForm(request.POST, prefix='analogue3')
        # analogues_image11_form = ImagesForm(request.POST, request.FILES, prefix='analogues_image11')
        # analogues_image21_form = ImagesForm(request.POST, request.FILES, prefix='analogues_image21')
        # analogues_image31_form = ImagesForm(request.POST, request.FILES, prefix='analogues_image31')
        # if analogues_image11_form.is_valid():
        #     item_analogues_image11 = analogues_image11_form.save()
        #     print('Сохранил11', item_analogues_image11.pk)
        #     if analogues_image21_form.is_valid():
        #         item_analogues_image21 = analogues_image21_form.save()
        #         print('Сохранил21')
        #         if analogues_image31_form.is_valid():
        #             item_analogues_image31 = analogues_image31_form.save()
        #             print('Сохранил31')

        if (form.is_valid() and client_person_data_form.is_valid() and object_of_assessment_form.is_valid() and analogues_form1.is_valid() and analogues_form2.is_valid() and analogues_form3.is_valid()
        and adjustment_form1.is_valid() and adjustment_form2.is_valid() and adjustment_form3.is_valid()):
    #         # <process form cleaned data>
            item_report = form.save()
            item_client = client_person_data_form.save()
            item_object = object_of_assessment_form.save()
            item_analogues1 = analogues_form1.save()
            item_analogues2 = analogues_form2.save()
            item_analogues3 = analogues_form3.save()
            item_adjustment1 = adjustment_form1.save()
            item_adjustment2 = adjustment_form2.save()
            item_adjustment3 = adjustment_form3.save()
            report = Report.objects.get(id=item_report.pk)
            analogue1 = Analogues.objects.get(id=item_analogues1.pk)
            analogue2 = Analogues.objects.get(id=item_analogues2.pk)
            analogue3 = Analogues.objects.get(id=item_analogues3.pk)
            report.clientPersonData_id = item_client.pk
            report.objectOfAssessment_id = item_object.pk
            report.analogue1_id = item_analogues1.pk
            report.analogue2_id = item_analogues2.pk
            report.analogue3_id = item_analogues3.pk
            report.save(update_fields=['clientPersonData_id', 'objectOfAssessment_id', 'analogue1_id', 'analogue2_id', 'analogue3_id',], force_update=True)
            analogue1.analogueAdjustments_id = item_adjustment1
            analogue1.save(update_fields=['analogueAdjustments_id', ], force_update=True)
            analogue2.analogueAdjustments_id = item_adjustment2
            analogue2.save(update_fields=['analogueAdjustments_id', ], force_update=True)
            analogue3.analogueAdjustments_id = item_adjustment3
            analogue3.save(update_fields=['analogueAdjustments_id', ], force_update=True)
            return redirect('notarius:update_item1', item_report.pk)

            #print("This is my newly created instance", Report.objects.get(contractNumber=self.request.POST.get('contractNumber')))
        return render(request, self.template_name, context={'form': form, "client_person_data_form": client_person_data_form, "object_of_assessment_form": object_of_assessment_form,
                                                            "analogues_form1": analogues_form1, "analogues_form2": analogues_form2, "analogues_form3": analogues_form3,
                                                            'adjustment_form1': adjustment_form1, 'adjustment_form2': adjustment_form2, 'adjustment_form3': adjustment_form3,
                                                            })

    # def post1(self, request, *args, **kwargs):
    #     print('1')
    #     form = self.form_class(request.POST, request.FILES)
    #     purpose_form = CreatePurposeOfAssessment(request.POST)
    #     person_data_form = CreatePersonDataForm(request.POST)
    #     print(self.request.POST.get('clientType'))
    #     if self.request.POST.get('clientType') == '2':
    #         print('лицо')
    #
    #         return render(request, "notarius/test.html", context={'form': form, "purpose_of_assessment": purpose_form, "client_person_data_form": person_data_form})
    #     else:
    #         print('не понятно')

    #
    #     else:
    #         print('физическое лицо')
    #         person_data_form = CreatePersonDataForm(request.POST)
    #         return render(request, self.template_name, context={'form': form, "purpose_of_assessment": purpose_form, "client_person_data_form": person_data_form})

        # if form.is_valid():
        #     i = form.cleaned_data['clientType']
        #     print(i)
        # else:
        #     i = form.cleaned_data['clientType']
        #     print(i)
        #     print('not valid')
        #     person_data_form = None
        #
        #     if (form.is_valid() and purpose_form.is_valid() and person_data_form == None):
        #         # <process form cleaned data>
        #         item_purpose = purpose_form.save()
        #         print(item_purpose.pk)
        #         #print("This is my newly created instance", Report.objects.get(contractNumber=self.request.POST.get('contractNumber')))
        #         return redirect(f"/notarius/")
        # return render(request, self.template_name, context={'form': form, "purpose_of_assessment": purpose_form})



#     def form_valid(self, form):
#         print('перед пурпос')
#         purpose_id = PurposeOfAssessment.objects.create(purposeOfAssessment1=self.request.POST.get('purposeOfAssessment1'))
#         print('после пурпос')
#         result = super().form_valid(form)
#         thisreport = Report.objects.get(id=self.object.pk)
#         thisreport.purposeOfAssessment_id = purpose_id.pk
# #        report = Report(purposeOfAssessment_id=purpose_id.pk)
#         thisreport.save(update_fields=['purposeOfAssessment_id', ], force_update=True)
#         print(self.request.POST.get('purposeOfAssessment1'))
#         print("This is my newly created instance", self.object.pk)
#
#         return result

# def add_item(request):
#
#     if request.method == "POST":
#
#         report_form = CreateReport(request.POST, request.FILES)
#         purpose_form = CreatePurposeOfAssessment(request.POST)
#
#         if (report_form.is_valid() and purpose_form.is_valid()):
#             report_form.save()
#             purpose_form.save()
#             thisreport = Report.pk
#             print(thisreport)
#             return redirect(f"/notarius/")
#
#     else:
#         report_form = CreateReport(request.POST, request.FILES)
#         purpose_form = CreatePurposeOfAssessment(request.POST)
#
#     data = {'report_form': report_form,
#                 'purpose_form': purpose_form}
#     return render(request, "notarius/additem.html", data)

# def add_item(request):
#     if request.method == "POST":
#         contractnum = request.POST.get("contractnum")
#         contractdate = request.POST.get("contractdate")
#         clientname = request.POST.get("clientname")
#         image = request.FILES['upload']
#         appraiser = request.user
#         item = Report(contractNumber=contractnum, conrtractDate=contractdate, clientname=clientname, image=image, appraiser=appraiser)
#         item.save()
#         my_id=item.pk
#         return redirect(f"/notarius/{my_id}")
#     return render(request, "notarius/additem.html")



def update_item(request, my_id):
    item = Report.objects.get(id=my_id)
    if request.method == "POST":
        item.contractNumber = request.POST.get("contractnum")
        item.conrtractDate = request.POST.get("contractdate")
        item.clientname = request.POST.get("clientname")
        item.image = request.FILES.get('upload', item.image)
        item.save()
    context = {'item': item}
    return render(request, "notarius/updateitem.html", context)


class UpdateReport(LoginRequiredMixin, UpdateView):
    model = Report
    fields = ['contractNumber', 'conrtractDate', 'reportNumber', 'dateOfAssessment', 'dateOfReport', 'documentsOfReport', 'purposeOfAssessment', 'clientName', 'objectTotalCost']
    context_object_name = 'item'
    template_name = 'notarius/additem.html'



    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context["client_person_data_form"] = CreatePersonDataForm(instance=ClientPersonData.objects.get(pk=self.object.clientPersonData_id))
        context["object_of_assessment_form"] = ObjectOfAssessmentForm(instance=ObjectOfAssessment.objects.get(pk=self.object.objectOfAssessment_id))
        context["analogues_form1"] = AnaloguesForm(instance=Analogues.objects.get(pk=self.object.analogue1_id), prefix='analogue1')
        context["analogues_form2"] = AnaloguesForm(instance=Analogues.objects.get(pk=self.object.analogue2_id), prefix='analogue2')
        context["analogues_form3"] = AnaloguesForm(instance=Analogues.objects.get(pk=self.object.analogue3_id), prefix='analogue3')
        context["adjustment_form1"] = AdjustmentsForm(instance=Adjustments.objects.get(pk=Analogues.objects.get(pk=self.object.analogue1_id).analogueAdjustments_id), prefix='analogue1')
        context["adjustment_form2"] = AdjustmentsForm(instance=Adjustments.objects.get(pk=Analogues.objects.get(pk=self.object.analogue2_id).analogueAdjustments_id), prefix='analogue2')
        context["adjustment_form3"] = AdjustmentsForm(instance=Adjustments.objects.get(pk=Analogues.objects.get(pk=self.object.analogue3_id).analogueAdjustments_id), prefix='analogue3')
        return context

    def post(self, *args, **kwargs):
        my_id = self.kwargs.get('pk')
        my_object = Report.objects.get(pk=my_id)
        form = CreateReport(self.request.POST, instance=my_object)
        client_person_data = ClientPersonData.objects.get(pk=my_object.clientPersonData_id)
        client_person_data_form = CreatePersonDataForm(self.request.POST, instance=client_person_data)
        object_of_assessment_form = ObjectOfAssessmentForm(self.request.POST, instance=ObjectOfAssessment.objects.get(pk=my_object.objectOfAssessment_id))
        analogues_form1 = AnaloguesForm(self.request.POST, prefix='analogue1', instance=Analogues.objects.get(pk=my_object.analogue1_id))
        analogues_form2 = AnaloguesForm(self.request.POST, prefix='analogue2', instance=Analogues.objects.get(pk=my_object.analogue2_id))
        analogues_form3 = AnaloguesForm(self.request.POST, prefix='analogue3', instance=Analogues.objects.get(pk=my_object.analogue3_id))
        adjustment_form1 = AdjustmentsForm(self.request.POST, prefix='analogue1', instance=Adjustments.objects.get(pk=Analogues.objects.get(pk=my_object.analogue1_id).analogueAdjustments_id))
        adjustment_form2 = AdjustmentsForm(self.request.POST, prefix='analogue2', instance=Adjustments.objects.get(pk=Analogues.objects.get(pk=my_object.analogue2_id).analogueAdjustments_id))
        adjustment_form3 = AdjustmentsForm(self.request.POST, prefix='analogue3', instance=Adjustments.objects.get(pk=Analogues.objects.get(pk=my_object.analogue3_id).analogueAdjustments_id))
        print('start validation')
        if (form.is_valid() and client_person_data_form.is_valid() and object_of_assessment_form.is_valid() and analogues_form1.is_valid() and analogues_form2.is_valid() and analogues_form3.is_valid()
        and adjustment_form1.is_valid() and adjustment_form2.is_valid() and adjustment_form3.is_valid()):
  #         # <process form cleaned data>
            my_object.save(update_fields=['contractNumber', 'conrtractDate', 'reportNumber', 'dateOfAssessment', 'dateOfReport', 'documentsOfReport', 'purposeOfAssessment', 'clientName', 'objectTotalCost' ], force_update=True)
            my_object.objectTotalCostWord = num2rus.converter(my_object.objectTotalCost)
            my_object.save(update_fields=['objectTotalCostWord', ], force_update=True)
            client_person_data_form.save()
            object_of_assessment_form.save()
            analogues_form1.save()
            analogues_form2.save()
            analogues_form3.save()
            adjustment_form1.save()
            adjustment_form2.save()
            adjustment_form3.save()
            if 'make_report' in self.request.POST:
                report = my_object
                MakeReport(report)
                print(open(f'media/documents/notarius/{report.pk}.docx', 'rb'))
                chunk_size = 8192
                response = StreamingHttpResponse(FileWrapper(open(f'media/documents/notarius/{report.pk}.docx', 'rb'), chunk_size),
                                                 content_type=mimetypes.guess_type(f'media/documents/notarius/{report.pk}.docx')[0])
                response['Content-Length'] = getsize(f'media/documents/notarius/{report.pk}.docx')
                response['Content-Disposition'] = "attachment; filename=%s" % f'{report.reportNumber}.docx'
                return response
            return redirect('notarius:update_item1', my_id)

        else:
            print('фигня')
            my_id = self.kwargs.get('pk')
            my_object = Report.objects.get(pk=my_id)
            form = CreateReport(instance=my_object)
            client_person_data_form = CreatePersonDataForm(self.request.POST)
            object_of_assessment_form = ObjectOfAssessmentForm(self.request.POST)
            analogues_form1 = AnaloguesForm(self.request.POST, prefix='analogue1')
            analogues_form2 = AnaloguesForm(self.request.POST, prefix='analogue2')
            analogues_form3 = AnaloguesForm(self.request.POST, prefix='analogue3')
            adjustment_form1 = AdjustmentsForm(self.request.POST, prefix='analogue1')
            adjustment_form2 = AdjustmentsForm(self.request.POST, prefix='analogue2')
            adjustment_form3 = AdjustmentsForm(self.request.POST, prefix='analogue3')

        return render(self.request, self.template_name, context={'title_dander': 'Проверьте данные', 'form': form, "client_person_data_form": client_person_data_form, "object_of_assessment_form": object_of_assessment_form,
                                                            "analogues_form1": analogues_form1, "analogues_form2": analogues_form2, "analogues_form3": analogues_form3,
                                                            'adjustment_form1': adjustment_form1, 'adjustment_form2': adjustment_form2, 'adjustment_form3': adjustment_form3,
                                                            })

#     def get_success_url(self, *args, **kwargs):
#         print('start get success url')
#         """Detect the submit button used and act accordingly"""
#         if 'make_report' in self.request.POST:
#             report = self.object
# #            url = reverse_lazy('notarius:update_item1', kwargs={'pk': report.pk})
#             MakeReport(report)
#             print('1')
#             print(redirect('http://127.0.0.1:8000/media/documents/notarius/129.docx'))
#
#         else:
#             print('no make report')
#             url = reverse_lazy('notarius:update_item1', kwargs={'pk': self.object.pk})
#             return url


# def update_item1(request, my_id):
#     item = Report.objects.get(id=my_id)
#     print(item.image.url)
#     if request.method == "POST":
#         form = UpdateReport(request.POST, request.FILES)
#         if form.is_valid():
#             print(form.cleaned_data)
#             form.save()
#             item = Report.objects.get(id=my_id)
#             print(item.image.url)
#             context = {'form': form,
#                        'item': item}
#             return render(request, "notarius/updateitem1.html", context)
#         else:
#             print('неверно')
#     else:
#         form = UpdateReport()
#     context = {'form': form,
#                'item': item}
#     return render(request, "notarius/updateitem1.html", context)


def delete_item(request, my_id):
    item = Report.objects.get(id=my_id)
    if request.method == "POST":
        item.delete()
        return redirect("/notarius/")
    context = {'item': item}
    return render(request, "notarius/deleteitem.html", context)

class ReportDeleteView(DeleteView):
    model = Report
    success_url = reverse_lazy("notarius:index")


def main(request):
   return render(request, "notarius/main.html")


